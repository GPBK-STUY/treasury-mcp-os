"""
TreasuryOS for Business Owners
Get financing-ready. Know your numbers. Walk into the bank with confidence.
Run:  streamlit run smb_app.py
"""

import streamlit as st
import pandas as pd
import os
import sys
import io
import re
import tempfile
from pathlib import Path
from datetime import datetime

_DIR = Path(__file__).resolve().parent
if str(_DIR) not in sys.path:
    sys.path.insert(0, str(_DIR))

from tools.aggregator import get_cash_position
from tools.forecaster import forecast_cash_position
from tools.working_capital import analyze_working_capital
from tools.covenant_monitor import monitor_debt_covenants
from tools.credit_parser import parse_credit_report
from tools.credit_assessor import assess_credit_position

# ─── Config ─────────────────────────────────────────────────
st.set_page_config(page_title="TreasuryOS", page_icon="T", layout="wide", initial_sidebar_state="expanded")

# ─── Design System ──────────────────────────────────────────
NAVY = "#0B1426"
DARK_NAVY = "#060E1A"
CARD_BG = "#111D32"
CARD_BORDER = "#1E3050"
ACCENT = "#3B82F6"
ACCENT_LIGHT = "#60A5FA"
TEXT_PRIMARY = "#F1F5F9"
TEXT_SECONDARY = "#94A3B8"
TEXT_MUTED = "#64748B"
GREEN = "#10B981"
YELLOW = "#F59E0B"
RED = "#EF4444"
SURFACE = "#0F1A2E"

st.markdown(f"""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
    .stApp {{ background: {NAVY}; font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif; }}
    .main .block-container {{ padding: 1.5rem 2rem 2rem; max-width: 1200px; }}
    #MainMenu, footer {{ visibility: hidden; }}

    section[data-testid="stSidebar"] {{
        background: {DARK_NAVY}; border-right: 1px solid {CARD_BORDER};
    }}
    section[data-testid="stSidebar"] .stMarkdown p,
    section[data-testid="stSidebar"] .stMarkdown li {{ color: {TEXT_SECONDARY}; font-size: 0.875rem; }}
    section[data-testid="stSidebar"] h1 {{ color: {TEXT_PRIMARY}; font-size: 1.25rem; font-weight: 600; letter-spacing: -0.01em; }}
    section[data-testid="stSidebar"] h3 {{ color: {TEXT_MUTED}; font-size: 0.7rem; font-weight: 600; text-transform: uppercase; letter-spacing: 0.08em; }}
    section[data-testid="stSidebar"] hr {{ border-color: {CARD_BORDER}; margin: 0.75rem 0; }}

    h1 {{ color: {TEXT_PRIMARY} !important; font-weight: 700 !important; font-size: 1.75rem !important; letter-spacing: -0.02em !important; margin-bottom: 0.25rem !important; }}
    h2 {{ color: {TEXT_PRIMARY} !important; font-weight: 600 !important; font-size: 1.25rem !important; letter-spacing: -0.01em !important; }}
    h3 {{ color: {TEXT_SECONDARY} !important; font-weight: 500 !important; font-size: 1rem !important; }}
    p, li {{ color: {TEXT_SECONDARY}; }}
    .stCaption p {{ color: {TEXT_MUTED} !important; font-size: 0.8rem !important; }}

    div[data-testid="stMetric"] {{
        background: {CARD_BG}; border: 1px solid {CARD_BORDER}; border-radius: 10px; padding: 1.25rem;
    }}
    div[data-testid="stMetric"] label {{ color: {TEXT_MUTED} !important; font-size: 0.75rem !important; font-weight: 500; text-transform: uppercase; letter-spacing: 0.05em; }}
    div[data-testid="stMetric"] [data-testid="stMetricValue"] {{ color: {TEXT_PRIMARY} !important; font-size: 1.5rem !important; font-weight: 700; }}
    div[data-testid="stMetric"] [data-testid="stMetricDelta"] {{ font-size: 0.8rem !important; }}

    .stDataFrame {{ border-radius: 10px; overflow: hidden; }}
    .stDataFrame [data-testid="stDataFrameResizable"] {{ background: {CARD_BG}; border: 1px solid {CARD_BORDER}; border-radius: 10px; }}

    .streamlit-expanderHeader {{
        background: {CARD_BG} !important; border: 1px solid {CARD_BORDER} !important;
        border-radius: 8px !important; color: {TEXT_PRIMARY} !important; font-weight: 500 !important; font-size: 0.9rem !important;
    }}
    .streamlit-expanderContent {{
        background: {SURFACE} !important; border: 1px solid {CARD_BORDER} !important;
        border-top: none !important; border-radius: 0 0 8px 8px !important;
    }}

    .stButton > button {{
        background: {ACCENT} !important; color: white !important; border: none !important;
        border-radius: 8px !important; font-weight: 600 !important; padding: 0.5rem 1.5rem !important; transition: all 0.2s !important;
    }}
    .stButton > button:hover {{ background: {ACCENT_LIGHT} !important; }}
    .stSlider label, .stNumberInput label, .stRadio label, .stFileUploader label {{ color: {TEXT_SECONDARY} !important; font-size: 0.85rem !important; }}
    .stRadio [data-testid="stMarkdownContainer"] p {{ color: {TEXT_SECONDARY} !important; }}
    .stAlert {{ border-radius: 8px; }}
    .stPlotlyChart, [data-testid="stArrowVegaLiteChart"] {{ background: {CARD_BG}; border-radius: 10px; padding: 0.5rem; border: 1px solid {CARD_BORDER}; }}

    .section-divider {{ border: none; border-top: 1px solid {CARD_BORDER}; margin: 1.5rem 0; }}
    .card {{ background: {CARD_BG}; border: 1px solid {CARD_BORDER}; border-radius: 10px; padding: 1.25rem; margin-bottom: 1rem; }}
    .card-header {{ color: {TEXT_MUTED}; font-size: 0.7rem; font-weight: 600; text-transform: uppercase; letter-spacing: 0.08em; margin-bottom: 0.75rem; }}
    .card-value {{ color: {TEXT_PRIMARY}; font-size: 1.75rem; font-weight: 700; line-height: 1.2; }}
    .card-subtitle {{ color: {TEXT_SECONDARY}; font-size: 0.8rem; margin-top: 0.25rem; }}
    .dot-green {{ color: {GREEN}; }}
    .dot-yellow {{ color: {YELLOW}; }}
    .dot-red {{ color: {RED}; }}
    .status-row {{ display: flex; align-items: center; gap: 0.5rem; padding: 0.625rem 0; border-bottom: 1px solid {CARD_BORDER}; }}
    .status-row:last-child {{ border-bottom: none; }}
    .status-dot {{ width: 8px; height: 8px; border-radius: 50%; display: inline-block; flex-shrink: 0; }}
    .status-dot.green {{ background: {GREEN}; box-shadow: 0 0 6px {GREEN}40; }}
    .status-dot.yellow {{ background: {YELLOW}; box-shadow: 0 0 6px {YELLOW}40; }}
    .status-dot.red {{ background: {RED}; box-shadow: 0 0 6px {RED}40; }}
    .status-label {{ color: {TEXT_PRIMARY}; font-size: 0.875rem; font-weight: 500; }}
    .status-detail {{ color: {TEXT_MUTED}; font-size: 0.8rem; margin-left: auto; }}
    .stat-grid {{ display: grid; grid-template-columns: 1fr 1fr; gap: 0.5rem; }}
    .stat-item {{ padding: 0.5rem 0; }}
    .stat-label {{ color: {TEXT_MUTED}; font-size: 0.75rem; }}
    .stat-value {{ color: {TEXT_PRIMARY}; font-size: 0.9rem; font-weight: 500; }}
    .tag {{ display: inline-block; padding: 0.2rem 0.6rem; border-radius: 4px; font-size: 0.75rem; font-weight: 600; }}
    .tag-green {{ background: {GREEN}20; color: {GREEN}; }}
    .tag-yellow {{ background: {YELLOW}20; color: {YELLOW}; }}
    .tag-red {{ background: {RED}20; color: {RED}; }}
    .top-bar {{ display: flex; justify-content: space-between; align-items: center; margin-bottom: 1.5rem; }}
    .top-bar-title {{ color: {TEXT_PRIMARY}; font-size: 1.75rem; font-weight: 700; letter-spacing: -0.02em; }}
    .top-bar-sub {{ color: {TEXT_MUTED}; font-size: 0.85rem; }}

    .step-num {{
        display: inline-flex; align-items: center; justify-content: center;
        width: 28px; height: 28px; border-radius: 50%; font-weight: 700; font-size: 0.8rem;
        flex-shrink: 0;
    }}
    .step-active {{ background: {ACCENT}; color: white; }}
    .step-done {{ background: {GREEN}; color: white; }}
    .step-pending {{ background: {CARD_BORDER}; color: {TEXT_MUTED}; }}
</style>
""", unsafe_allow_html=True)


# ─── Session State ──────────────────────────────────────────
if "data_dir" not in st.session_state:
    st.session_state.data_dir = str(_DIR / "sample_data")
    st.session_state.using_sample = True
if "uploaded_files" not in st.session_state:
    st.session_state.uploaded_files = {}


# ─── Helpers ────────────────────────────────────────────────
def fmt(value, decimals=0):
    if value is None: return "—"
    if abs(value) >= 1_000_000: return f"${value/1_000_000:,.{decimals}f}M"
    if abs(value) >= 1_000: return f"${value/1_000:,.{decimals}f}K"
    return f"${value:,.{decimals}f}"

def dot_color(status):
    s = status.lower()
    if s in ("compliant","strong","healthy","excellent","good","positive","low_risk"): return "green"
    if s in ("warning","fair","moderate","marginal","medium_risk","needs_work"): return "yellow"
    return "red"

def status_html(status, label=None, detail=None):
    c = dot_color(status)
    lbl = label or status.replace("_"," ").title()
    d = f'<span class="status-detail">{detail}</span>' if detail else ""
    return f'<div class="status-row"><span class="status-dot {c}"></span><span class="status-label">{lbl}</span>{d}</div>'

def tag_html(status, text=None):
    c = dot_color(status)
    t = text or status.replace("_"," ").title()
    return f'<span class="tag tag-{c}">{t}</span>'

def metric_card(label, value, subtitle=""):
    sub = f'<div class="card-subtitle">{subtitle}</div>' if subtitle else ""
    return f'<div class="card"><div class="card-header">{label}</div><div class="card-value">{value}</div>{sub}</div>'

def run_safe(fn, *a, **kw):
    try: return fn(*a, **kw), None
    except FileNotFoundError as e: return None, f"Missing: {e}"
    except Exception as e: return None, str(e)

def save_uploads(files):
    d = tempfile.mkdtemp(prefix="treasuryos_")
    for name, f in files.items():
        if hasattr(f, 'getbuffer'):
            with open(os.path.join(d, name), "wb") as out:
                out.write(f.getbuffer())
        elif isinstance(f, bytes):
            with open(os.path.join(d, name), "wb") as out:
                out.write(f)
    return d


# ─── File Type Auto-Detection ──────────────────────────────
_FILE_MAP_KEYWORDS = {
    "accounts.csv":          ["account", "balance", "bank", "checking", "savings", "deposit"],
    "transactions.csv":      ["transaction", "debit", "credit", "payment", "transfer", "date", "amount"],
    "vendors.csv":           ["vendor", "supplier", "payable", "invoice", "due_date", "discount"],
    "covenants.csv":         ["covenant", "ratio", "threshold", "compliance", "facility", "debt"],
    "fx_rates.csv":          ["currency", "exchange", "rate", "fx", "eur", "gbp", "jpy", "usd"],
    "personal_credit.csv":   ["fico", "credit_score", "utilization", "derogatory", "late_payment", "personal"],
    "business_credit.csv":   ["paydex", "dbt", "lien", "judgment", "years_in_business", "business"],
}

def _guess_csv_type(df, original_name=""):
    cols_lower = {c.lower().replace(" ", "_") for c in df.columns}
    name_lower = original_name.lower()
    best_match, best_score = None, 0
    for target_csv, keywords in _FILE_MAP_KEYWORDS.items():
        score = sum(2 for kw in keywords if any(kw in col for col in cols_lower))
        target_base = target_csv.replace(".csv", "")
        if target_base in name_lower:
            score += 5
        if score > best_score:
            best_score = score
            best_match = target_csv
    return best_match if best_score >= 2 else None

def _parse_excel(file_bytes, filename):
    results = {}
    try:
        import openpyxl
        xls = pd.ExcelFile(io.BytesIO(file_bytes), engine="openpyxl")
        for sheet in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet)
            if df.empty: continue
            csv_type = _guess_csv_type(df, f"{filename}_{sheet}")
            if csv_type and csv_type not in results:
                buf = io.BytesIO(); df.to_csv(buf, index=False)
                results[csv_type] = buf.getvalue()
        if not results and len(xls.sheet_names) == 1:
            df = pd.read_excel(xls, sheet_name=0)
            csv_type = _guess_csv_type(df, filename)
            if csv_type:
                buf = io.BytesIO(); df.to_csv(buf, index=False)
                results[csv_type] = buf.getvalue()
    except Exception: pass
    return results

def _extract_pdf_text(file_bytes):
    """Extract all text from a PDF."""
    text = ""
    try:
        import pdfplumber
        pdf = pdfplumber.open(io.BytesIO(file_bytes))
        for pg in pdf.pages:
            pg_text = pg.extract_text()
            if pg_text:
                text += pg_text + "\n"
        pdf.close()
    except Exception:
        pass
    return text


def _parse_text_to_credit(text, filename):
    """Try to extract credit data from raw text using pattern matching."""
    results = {}
    text_lower = text.lower()

    # ── Personal Credit Report Detection ──
    personal_indicators = ["fico", "credit score", "transunion", "equifax", "experian",
                           "payment history", "revolving", "utilization", "inquiries",
                           "tradeline", "credit report", "derogatory", "collections",
                           "credit karma", "annual credit report", "vantagescore"]
    personal_hits = sum(1 for k in personal_indicators if k in text_lower)

    if personal_hits >= 3:
        row = {
            "borrower_name": "", "ssn_last4": "", "credit_score": 0, "score_model": "FICO 8",
            "report_date": "", "total_tradelines": 0, "open_tradelines": 0,
            "derogatory_marks": 0, "collections": 0, "public_records": 0,
            "total_revolving_balance": 0, "total_revolving_limit": 0,
            "revolving_utilization_pct": 0, "total_installment_balance": 0,
            "monthly_installment_payments": 0, "oldest_account_years": 0,
            "recent_inquiries_6mo": 0, "late_payments_30d": 0, "late_payments_60d": 0,
            "late_payments_90d": 0, "payment_history_pct": 0, "bankruptcies": 0,
            "foreclosures": 0, "tax_liens": 0,
        }

        # Extract FICO / credit score
        import re
        score_patterns = [
            r'(?:fico|credit\s*score|vantage\s*score|score)[:\s]*(\d{3})',
            r'(\d{3})\s*(?:fico|credit\s*score|vantage)',
            r'score[:\s]*(\d{3})',
        ]
        for pat in score_patterns:
            m = re.search(pat, text_lower)
            if m:
                val = int(m.group(1))
                if 300 <= val <= 850:
                    row["credit_score"] = val
                    break

        # Extract utilization
        util_patterns = [
            r'(?:utilization|usage)[:\s]*([\d.]+)\s*%',
            r'([\d.]+)\s*%\s*(?:utilization|usage)',
        ]
        for pat in util_patterns:
            m = re.search(pat, text_lower)
            if m:
                row["revolving_utilization_pct"] = float(m.group(1))
                break

        # Extract payment history percentage
        hist_patterns = [r'(?:payment\s*history|on.time)[:\s]*([\d.]+)\s*%']
        for pat in hist_patterns:
            m = re.search(pat, text_lower)
            if m:
                row["payment_history_pct"] = float(m.group(1))
                break

        # Extract tradeline counts
        trade_patterns = [r'(\d+)\s*(?:total\s*)?(?:accounts?|tradelines?)']
        for pat in trade_patterns:
            m = re.search(pat, text_lower)
            if m:
                row["total_tradelines"] = int(m.group(1))
                break

        open_patterns = [r'(\d+)\s*open\s*(?:accounts?|tradelines?)']
        for pat in open_patterns:
            m = re.search(pat, text_lower)
            if m:
                row["open_tradelines"] = int(m.group(1))
                break

        # Extract late payments
        late30 = re.search(r'(\d+)\s*(?:late|past\s*due).*?30', text_lower)
        late60 = re.search(r'(\d+)\s*(?:late|past\s*due).*?60', text_lower)
        late90 = re.search(r'(\d+)\s*(?:late|past\s*due).*?90', text_lower)
        if late30: row["late_payments_30d"] = int(late30.group(1))
        if late60: row["late_payments_60d"] = int(late60.group(1))
        if late90: row["late_payments_90d"] = int(late90.group(1))

        # Derogatory marks
        derog = re.search(r'(\d+)\s*derogator', text_lower)
        if derog: row["derogatory_marks"] = int(derog.group(1))

        # Collections
        coll = re.search(r'(\d+)\s*collection', text_lower)
        if coll: row["collections"] = int(coll.group(1))

        # Inquiries
        inq = re.search(r'(\d+)\s*(?:inquir|hard\s*pull)', text_lower)
        if inq: row["recent_inquiries_6mo"] = int(inq.group(1))

        # Bankruptcies
        bk = re.search(r'(\d+)\s*bankruptc', text_lower)
        if bk: row["bankruptcies"] = int(bk.group(1))

        # Balance extraction
        bal_patterns = [
            (r'revolving\s*balance[:\s]*\$?([\d,]+)', "total_revolving_balance"),
            (r'revolving\s*(?:credit\s*)?limit[:\s]*\$?([\d,]+)', "total_revolving_limit"),
            (r'installment\s*balance[:\s]*\$?([\d,]+)', "total_installment_balance"),
        ]
        for pat, field in bal_patterns:
            m = re.search(pat, text_lower)
            if m:
                row[field] = float(m.group(1).replace(",", ""))

        if row["credit_score"] > 0:
            df = pd.DataFrame([row])
            buf = io.BytesIO(); df.to_csv(buf, index=False)
            results["personal_credit.csv"] = buf.getvalue()

    # ── Business Credit Report Detection ──
    business_indicators = ["paydex", "duns", "d&b", "dun & bradstreet", "experian business",
                           "intelliscore", "business credit", "trade experiences",
                           "days beyond terms", "payment trend", "ucc filing",
                           "sic code", "naics", "nav.com"]
    business_hits = sum(1 for k in business_indicators if k in text_lower)

    if business_hits >= 2:
        row = {
            "business_name": "", "duns_number": "", "report_date": "",
            "paydex_score": 0, "intelliscore": 0, "years_in_business": 0,
            "sic_code": "", "industry": "", "total_trade_experiences": 0,
            "current_pct": 0, "days_beyond_terms_avg": 0, "high_credit": 0,
            "total_balance_outstanding": 0, "payment_trend": "stable",
            "derogatory_count": 0, "liens": 0, "judgments": 0, "ucc_filings": 0,
            "bankruptcy_flag": "false", "d_and_b_rating": "", "annual_revenue": 0,
            "employees": 0, "trade_payment_index": 0,
        }

        import re
        paydex = re.search(r'paydex[:\s]*(\d{1,3})', text_lower)
        if paydex:
            val = int(paydex.group(1))
            if 0 <= val <= 100: row["paydex_score"] = val

        yrs = re.search(r'(\d+)\s*(?:years?\s*in\s*business|years?\s*established|yrs)', text_lower)
        if yrs: row["years_in_business"] = int(yrs.group(1))

        dbt = re.search(r'(?:days?\s*beyond\s*terms?|dbt)[:\s]*([\d.]+)', text_lower)
        if dbt: row["days_beyond_terms_avg"] = float(dbt.group(1))

        trade_exp = re.search(r'(\d+)\s*trade\s*experience', text_lower)
        if trade_exp: row["total_trade_experiences"] = int(trade_exp.group(1))

        liens_m = re.search(r'(\d+)\s*lien', text_lower)
        if liens_m: row["liens"] = int(liens_m.group(1))

        judgments_m = re.search(r'(\d+)\s*judgment', text_lower)
        if judgments_m: row["judgments"] = int(judgments_m.group(1))

        ucc_m = re.search(r'(\d+)\s*ucc', text_lower)
        if ucc_m: row["ucc_filings"] = int(ucc_m.group(1))

        if row["paydex_score"] > 0 or row["years_in_business"] > 0:
            df = pd.DataFrame([row])
            buf = io.BytesIO(); df.to_csv(buf, index=False)
            results["business_credit.csv"] = buf.getvalue()

    # ── Bank Statement Detection ──
    bank_indicators = ["checking", "savings", "account number", "routing",
                       "beginning balance", "ending balance", "deposits",
                       "withdrawals", "statement period", "bank statement",
                       "available balance", "current balance", "account summary"]
    bank_hits = sum(1 for k in bank_indicators if k in text_lower)

    if bank_hits >= 2:
        import re
        accounts = []
        # Try to find balances
        balance_patterns = [
            r'(?:ending|current|available|closing)\s*balance[:\s]*\$?([\d,]+\.?\d*)',
            r'\$\s*([\d,]+\.?\d{2})\s*(?:ending|current|available|closing)',
        ]
        for pat in balance_patterns:
            for m in re.finditer(pat, text_lower):
                bal = float(m.group(1).replace(",", ""))
                if bal > 0:
                    accounts.append({
                        "account_id": f"ACCT-PDF-{len(accounts)+1}",
                        "bank_name": "From Statement",
                        "account_type": "checking",
                        "currency": "USD",
                        "balance": bal,
                        "yield_rate_bps": 0,
                        "last_updated": "",
                    })

        # If no specific balances found, try to find any large dollar amounts
        if not accounts:
            dollar_amounts = re.findall(r'\$\s*([\d,]+\.?\d{2})', text)
            amounts = [float(a.replace(",", "")) for a in dollar_amounts if float(a.replace(",", "")) > 1000]
            if amounts:
                # Use the largest as the balance (likely ending balance)
                bal = max(amounts)
                accounts.append({
                    "account_id": "ACCT-PDF-1",
                    "bank_name": "From Statement",
                    "account_type": "checking",
                    "currency": "USD",
                    "balance": bal,
                    "yield_rate_bps": 0,
                    "last_updated": "",
                })

        if accounts:
            df = pd.DataFrame(accounts)
            buf = io.BytesIO(); df.to_csv(buf, index=False)
            results["accounts.csv"] = buf.getvalue()

    return results


def _parse_pdf(file_bytes, filename):
    results = {}
    try:
        import pdfplumber

        # First try: extract structured tables
        pdf = pdfplumber.open(io.BytesIO(file_bytes))
        for i, pg in enumerate(pdf.pages):
            for j, table in enumerate(pg.extract_tables()):
                if not table or len(table) < 2: continue
                df = pd.DataFrame(table[1:], columns=table[0])
                csv_type = _guess_csv_type(df, f"{filename}_p{i}_t{j}")
                if csv_type and csv_type not in results:
                    buf = io.BytesIO(); df.to_csv(buf, index=False)
                    results[csv_type] = buf.getvalue()
        pdf.close()

        # Second try: if no tables found, extract text and parse with patterns
        if not results:
            text = _extract_pdf_text(file_bytes)
            if text.strip():
                text_results = _parse_text_to_credit(text, filename)
                for k, v in text_results.items():
                    if k not in results:
                        results[k] = v

    except Exception: pass
    return results

def _parse_docx(file_bytes, filename):
    results = {}
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))

        # First try: extract tables
        for j, table in enumerate(doc.tables):
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if len(rows) < 2: continue
            df = pd.DataFrame(rows[1:], columns=rows[0])
            csv_type = _guess_csv_type(df, f"{filename}_t{j}")
            if csv_type and csv_type not in results:
                buf = io.BytesIO(); df.to_csv(buf, index=False)
                results[csv_type] = buf.getvalue()

        # Second try: if no tables found, extract text and parse with patterns
        if not results:
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            if text.strip():
                text_results = _parse_text_to_credit(text, filename)
                for k, v in text_results.items():
                    if k not in results:
                        results[k] = v
    except Exception: pass
    return results

def parse_uploaded_files(uploaded_files):
    all_csvs = {}
    for uf in uploaded_files:
        fname = uf.name.lower()
        raw = uf.read(); uf.seek(0)
        if fname.endswith(".csv"):
            try:
                df = pd.read_csv(io.BytesIO(raw))
                csv_type = _guess_csv_type(df, uf.name)
                if csv_type and csv_type not in all_csvs:
                    all_csvs[csv_type] = raw
                elif csv_type is None:
                    safe_name = re.sub(r'[^a-z0-9_.]', '_', fname)
                    all_csvs[safe_name] = raw
            except Exception: pass
        elif fname.endswith((".xlsx", ".xls")):
            for k, v in _parse_excel(raw, uf.name).items():
                if k not in all_csvs: all_csvs[k] = v
        elif fname.endswith(".pdf"):
            for k, v in _parse_pdf(raw, uf.name).items():
                if k not in all_csvs: all_csvs[k] = v
        elif fname.endswith(".docx"):
            for k, v in _parse_docx(raw, uf.name).items():
                if k not in all_csvs: all_csvs[k] = v
    return all_csvs


# ─── Sidebar ────────────────────────────────────────────────
with st.sidebar:
    st.markdown("# T  TreasuryOS")
    st.caption("For Business Owners")
    st.markdown("---")

    _nav_options = ["Home", "Get Your Score", "My Cash", "My Credit", "Business Health"]
    _default_idx = _nav_options.index(st.session_state.get("nav_page", "Home")) if st.session_state.get("nav_page", "Home") in _nav_options else 0
    page = st.radio("nav", _nav_options, index=_default_idx, label_visibility="collapsed")
    st.session_state.nav_page = page

    st.markdown("---")

    # ── Data source ──
    st.markdown("### Your Data")
    src = st.radio("src", ["Try Demo", "Upload My Files"], label_visibility="collapsed")

    if src == "Upload My Files":
        st.session_state.using_sample = False

        st.markdown(f"""<p style="color:{TEXT_SECONDARY};font-size:0.8rem;margin-bottom:0.25rem;">
        Upload these for the best score (most important first):</p>""", unsafe_allow_html=True)

        upload_guide = [
            ("1. Personal credit report", "Pull free from annualcreditreport.com or Credit Karma"),
            ("2. Business credit report", "From Nav.com, D&B, or Experian Business"),
            ("3. Bank statements", "Download from your bank as CSV or PDF"),
            ("4. Existing loan docs", "Any current loan agreements or covenant schedules"),
        ]
        for doc, source in upload_guide:
            st.markdown(f'<div style="padding:0.3rem 0;border-bottom:1px solid {CARD_BORDER};">'
                        f'<span style="color:{TEXT_PRIMARY};font-size:0.8rem;font-weight:600;">{doc}</span>'
                        f'<br><span style="color:{TEXT_MUTED};font-size:0.7rem;">{source}</span></div>',
                        unsafe_allow_html=True)

        st.markdown(f'<p style="color:{TEXT_MUTED};font-size:0.7rem;margin-top:0.5rem;">+ Enter your working capital numbers on the score page for best results</p>', unsafe_allow_html=True)
        st.caption("We accept CSV, Excel, Word, and PDF files.")

        uploaded_files = st.file_uploader(
            "Upload files",
            type=["csv", "xlsx", "xls", "docx", "pdf"],
            accept_multiple_files=True,
            key="multi_upload",
            label_visibility="collapsed"
        )
        if uploaded_files:
            parsed = parse_uploaded_files(uploaded_files)
            if parsed:
                st.session_state.data_dir = save_uploads(parsed)
                st.session_state.uploaded_files = parsed
                st.success(f"{len(parsed)} file(s) loaded")
                for fname in parsed:
                    st.caption(f"  {fname}")
            else:
                st.warning(
                    "We couldn't pull numbers from those files automatically. "
                    "This can happen with scanned documents or image-based PDFs. "
                    "Try downloading a CSV or Excel export from your bank's website, "
                    "or from Credit Karma / Nav.com if it's a credit report."
                )
    else:
        st.session_state.using_sample = True
        st.session_state.data_dir = str(_DIR / "sample_data")

    st.markdown("---")
    st.markdown(f'<p style="color:{TEXT_MUTED}; font-size:0.7rem; text-align:center;">TreasuryOS v1.0</p>', unsafe_allow_html=True)


# ─── Main ───────────────────────────────────────────────────
DATA = st.session_state.data_dir
NOW = datetime.now().strftime("%b %d, %Y")


# ═══════════════════════════════════════════════════════════
#  HOME — LANDING PAGE
# ═══════════════════════════════════════════════════════════
if page == "Home":

    # ── Hero ──
    st.markdown(f"""
    <div style="text-align:center;padding:2.5rem 1rem 1.5rem;">
        <div style="font-size:2.5rem;font-weight:800;color:{TEXT_PRIMARY};line-height:1.2;letter-spacing:-0.03em;">
            Are you ready for<br>a business loan?
        </div>
        <div style="font-size:1.1rem;color:{TEXT_SECONDARY};margin-top:1rem;max-width:550px;margin-left:auto;margin-right:auto;">
            Find out in 5 minutes. See exactly what a bank will look at,
            what's helping you, and what to fix before you apply.
        </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown('<hr class="section-divider">', unsafe_allow_html=True)

    # ── What you'll learn ──
    st.markdown(f"""
    <div style="text-align:center;margin-bottom:1.5rem;">
        <span style="color:{TEXT_MUTED};font-size:0.75rem;font-weight:600;text-transform:uppercase;letter-spacing:0.08em;">What you'll get</span>
    </div>
    """, unsafe_allow_html=True)

    c1, c2, c3 = st.columns(3)
    with c1:
        st.markdown(f"""<div class="card" style="text-align:center;min-height:180px;">
            <div style="font-size:2rem;margin-bottom:0.5rem;">📊</div>
            <div style="color:{TEXT_PRIMARY};font-weight:600;font-size:0.95rem;margin-bottom:0.5rem;">Your Readiness Score</div>
            <div style="color:{TEXT_SECONDARY};font-size:0.85rem;">A clear number from 0-100 showing how likely you are to get approved.</div>
        </div>""", unsafe_allow_html=True)
    with c2:
        st.markdown(f"""<div class="card" style="text-align:center;min-height:180px;">
            <div style="font-size:2rem;margin-bottom:0.5rem;">🎯</div>
            <div style="color:{TEXT_PRIMARY};font-weight:600;font-size:0.95rem;margin-bottom:0.5rem;">What to Fix First</div>
            <div style="color:{TEXT_SECONDARY};font-size:0.85rem;">A prioritized list of exactly what to improve before you walk into the bank.</div>
        </div>""", unsafe_allow_html=True)
    with c3:
        st.markdown(f"""<div class="card" style="text-align:center;min-height:180px;">
            <div style="font-size:2rem;margin-bottom:0.5rem;">💰</div>
            <div style="color:{TEXT_PRIMARY};font-weight:600;font-size:0.95rem;margin-bottom:0.5rem;">Loans You Qualify For</div>
            <div style="color:{TEXT_SECONDARY};font-size:0.85rem;">See which loan products match your profile — SBA, line of credit, equipment financing, and more.</div>
        </div>""", unsafe_allow_html=True)

    st.markdown('<hr class="section-divider">', unsafe_allow_html=True)

    # ── How it works ──
    st.markdown(f"""
    <div style="text-align:center;margin-bottom:1.5rem;">
        <span style="color:{TEXT_MUTED};font-size:0.75rem;font-weight:600;text-transform:uppercase;letter-spacing:0.08em;">How it works</span>
    </div>
    """, unsafe_allow_html=True)

    s1, s2, s3 = st.columns(3)
    with s1:
        st.markdown(f"""<div class="card">
            <div style="color:{ACCENT};font-weight:700;font-size:1.1rem;margin-bottom:0.5rem;">Step 1</div>
            <div style="color:{TEXT_PRIMARY};font-weight:600;margin-bottom:0.25rem;">Upload your documents</div>
            <div style="color:{TEXT_SECONDARY};font-size:0.85rem;">Bank statements, credit reports — whatever you have. We accept PDF, Excel, Word, and CSV files.</div>
        </div>""", unsafe_allow_html=True)
    with s2:
        st.markdown(f"""<div class="card">
            <div style="color:{ACCENT};font-weight:700;font-size:1.1rem;margin-bottom:0.5rem;">Step 2</div>
            <div style="color:{TEXT_PRIMARY};font-weight:600;margin-bottom:0.25rem;">We analyze everything</div>
            <div style="color:{TEXT_SECONDARY};font-size:0.85rem;">Our engine checks the same things a bank does — your credit, cash, business history, and existing debt.</div>
        </div>""", unsafe_allow_html=True)
    with s3:
        st.markdown(f"""<div class="card">
            <div style="color:{ACCENT};font-weight:700;font-size:1.1rem;margin-bottom:0.5rem;">Step 3</div>
            <div style="color:{TEXT_PRIMARY};font-weight:600;margin-bottom:0.25rem;">Get your game plan</div>
            <div style="color:{TEXT_SECONDARY};font-size:0.85rem;">See your score, what to fix, which loans you qualify for, and what documents to bring to the bank.</div>
        </div>""", unsafe_allow_html=True)

    st.markdown('<hr class="section-divider">', unsafe_allow_html=True)

    # ── CTA ──
    st.markdown(f"""
    <div style="text-align:center;padding:1rem 0 2rem;">
        <div style="color:{TEXT_SECONDARY};font-size:0.95rem;margin-bottom:1rem;">
            Ready? Click <b style="color:{TEXT_PRIMARY};">Get Your Score</b> in the sidebar to start.<br>
            Or try the demo first to see how it works.
        </div>
    </div>
    """, unsafe_allow_html=True)

    # Quick-start buttons
    col_l, col_c, col_r = st.columns([1, 2, 1])
    with col_c:
        if st.button("See a Demo Score", type="primary", use_container_width=True):
            st.session_state.using_sample = True
            st.session_state.data_dir = str(_DIR / "sample_data")
            st.session_state.nav_page = "Get Your Score"
            st.rerun()


# ═══════════════════════════════════════════════════════════
#  PAGE 2: GET YOUR SCORE (Financing Readiness)
# ═══════════════════════════════════════════════════════════
elif page == "Get Your Score":
    st.markdown(f'<div class="top-bar"><div><div class="top-bar-title">Get Your Financing Score</div><div class="top-bar-sub">See how a bank will evaluate your business &middot; {NOW}</div></div></div>', unsafe_allow_html=True)

    # ── Onboarding for demo users ──
    if st.session_state.get("using_sample", True):
        st.info(
            "**You're viewing a demo.** To see your real score, switch to "
            "**Upload My Files** in the sidebar and drop in your bank statements "
            "and credit reports. We accept CSV, Excel, Word, and PDF."
        )

    # ── Collect all data ──
    cash, cash_e = run_safe(get_cash_position, DATA)
    cov, cov_e = run_safe(monitor_debt_covenants, DATA)
    credit, credit_e = run_safe(parse_credit_report, DATA)
    assessment, assess_e = run_safe(assess_credit_position, DATA)

    # ── Optional working capital inputs ──
    with st.expander("Optional: Enter your working capital numbers for a more accurate score"):
        st.caption("You can find these on your balance sheet and P&L statement.")
        wc_c1, wc_c2 = st.columns(2)
        with wc_c1:
            fr_current_assets = st.number_input("Current Assets ($)", min_value=0, value=0, step=100_000, key="fr_ca")
            fr_accounts_receivable = st.number_input("Accounts Receivable ($)", min_value=0, value=0, step=100_000, key="fr_ar")
            fr_annual_revenue = st.number_input("Annual Revenue ($)", min_value=0, value=0, step=500_000, key="fr_rev")
        with wc_c2:
            fr_current_liabilities = st.number_input("Current Liabilities ($)", min_value=0, value=0, step=100_000, key="fr_cl")
            fr_accounts_payable = st.number_input("Accounts Payable ($)", min_value=0, value=0, step=100_000, key="fr_ap")
            fr_annual_cogs = st.number_input("Annual COGS ($)", min_value=0, value=0, step=500_000, key="fr_cogs")

    wc_data = None
    if fr_current_assets > 0 and fr_current_liabilities > 0 and fr_annual_revenue > 0:
        wc_data, _ = run_safe(analyze_working_capital,
                              float(fr_current_assets), float(fr_current_liabilities),
                              float(fr_annual_revenue), float(fr_accounts_receivable),
                              float(fr_accounts_payable), float(fr_annual_cogs))

    # ════════════════════════════════════════════════════════
    #  SCORING ENGINE (6 weighted categories)
    # ════════════════════════════════════════════════════════
    scores = {}

    # ── 1. Personal Credit (25%) ──
    pc_score = 0
    pc_findings = []
    pc_actions = []
    if credit and credit.get("personal_profiles"):
        p = credit["personal_profiles"][0]
        fico = p.get("credit_score", 0)
        util = p.get("revolving_utilization_pct", 100)
        derogs = p.get("derogatory_marks", 0)
        late30 = p.get("late_payments_30d", 0)
        late60 = p.get("late_payments_60d", 0)
        late90 = p.get("late_payments_90d", 0)

        if fico >= 750: pc_score += 40; pc_findings.append(f"FICO {fico} — excellent, you'll qualify for the best rates")
        elif fico >= 700: pc_score += 30; pc_findings.append(f"FICO {fico} — good, you qualify for most bank loans")
        elif fico >= 680: pc_score += 20; pc_findings.append(f"FICO {fico} — okay, but expect higher interest rates")
        elif fico >= 640: pc_score += 10; pc_findings.append(f"FICO {fico} — below what most banks want (680+)"); pc_actions.append("Work on raising your FICO above 680 — pay down credit cards and dispute any errors on your report")
        else: pc_findings.append(f"FICO {fico} — this will limit you to SBA microloans or secured lending"); pc_actions.append("Hold off on applying until your FICO improves above 640")

        if util <= 20: pc_score += 25; pc_findings.append(f"Credit card usage {util:.0f}% — excellent")
        elif util <= 30: pc_score += 20; pc_findings.append(f"Credit card usage {util:.0f}% — good")
        elif util <= 50: pc_score += 10; pc_findings.append(f"Credit card usage {util:.0f}% — banks want this under 30%"); pc_actions.append(f"Pay down your credit card balances to get usage under 30% (currently {util:.0f}%)")
        else: pc_findings.append(f"Credit card usage {util:.0f}% — too high, will hurt your application"); pc_actions.append(f"Get your credit card usage from {util:.0f}% down to under 30% before applying")

        total_lates = late30 + late60 + late90
        if total_lates == 0 and derogs == 0: pc_score += 25; pc_findings.append("Clean payment history — no late payments or negative marks")
        elif total_lates <= 1 and derogs == 0: pc_score += 15; pc_findings.append(f"{total_lates} late payment on your record")
        else:
            if total_lates > 0: pc_findings.append(f"{total_lates} late payments on your record")
            if derogs > 0: pc_findings.append(f"{derogs} negative mark(s) on your credit"); pc_actions.append("Try to get negative marks removed — call the creditor and negotiate a pay-for-delete agreement")

        inquiries = p.get("recent_inquiries_6mo", 0)
        if inquiries <= 2: pc_score += 10
        elif inquiries <= 4: pc_score += 5; pc_findings.append(f"{inquiries} recent credit checks — banks may wonder if you're desperate for credit")
        else: pc_findings.append(f"{inquiries} recent credit checks — stop applying for credit until you're ready"); pc_actions.append("Don't apply for any new credit for 6 months before your loan application")
    else:
        pc_findings.append("No personal credit data uploaded yet")
        pc_actions.append("Upload your credit report so we can give you a complete readiness score")

    scores["Your Personal Credit"] = {"score": pc_score, "weight": 0.25, "findings": pc_findings, "actions": pc_actions}

    # ── 2. Business Credit (20%) ──
    bc_score = 0
    bc_findings = []
    bc_actions = []
    if credit and credit.get("business_profile"):
        bp = credit["business_profile"]
        paydex = bp.get("paydex_score", 0)
        yrs = bp.get("years_in_business", 0)
        dbt = bp.get("days_beyond_terms_avg", 99)
        liens = bp.get("liens", 0)
        judgments = bp.get("judgments", 0)

        if paydex >= 80: bc_score += 35; bc_findings.append(f"Paydex {paydex} — strong, means you pay vendors on time or early")
        elif paydex >= 65: bc_score += 20; bc_findings.append(f"Paydex {paydex} — okay, but banks like 80+"); bc_actions.append("Pay your vendors faster to push your Paydex above 80")
        elif paydex >= 50: bc_score += 10; bc_findings.append(f"Paydex {paydex} — shows slow payment to vendors"); bc_actions.append(f"Your Paydex of {paydex} tells banks you pay late — make paying vendors on time a priority")
        else: bc_findings.append(f"Paydex {paydex} — this is a red flag for lenders"); bc_actions.append("Focus on paying every vendor on time for the next 6 months to rebuild your Paydex score")

        if yrs >= 5: bc_score += 25; bc_findings.append(f"{yrs} years in business — established track record")
        elif yrs >= 2: bc_score += 15; bc_findings.append(f"{yrs} years in business — meets minimum for most bank loans")
        elif yrs >= 1: bc_score += 5; bc_findings.append(f"Only {yrs} year(s) in business — many lenders want 2+"); bc_actions.append("Consider SBA microloans which work with newer businesses")
        else: bc_findings.append("Less than 1 year in business — most banks won't lend yet"); bc_actions.append("Build more operating history — look into SBA microloans or community development lenders")

        if dbt <= 10: bc_score += 20; bc_findings.append(f"Paying vendors {dbt:.0f} days beyond terms — on time")
        elif dbt <= 30: bc_score += 10; bc_findings.append(f"Paying vendors {dbt:.0f} days late on average"); bc_actions.append("Catch up on vendor payments — being 30+ days late shows up on your business credit")
        else: bc_findings.append(f"Paying vendors {dbt:.0f} days late — this will concern lenders"); bc_actions.append(f"You're paying {dbt:.0f} days late on average — make this your top priority to fix")

        if liens > 0: bc_score = max(bc_score - 15, 0); bc_findings.append(f"{liens} tax lien(s) on file — banks check for these"); bc_actions.append("Resolve tax liens before applying — this is often a deal-breaker")
        if judgments > 0: bc_score = max(bc_score - 15, 0); bc_findings.append(f"{judgments} judgment(s) against your business"); bc_actions.append("Clear any legal judgments before approaching a lender")
    else:
        bc_findings.append("No business credit data uploaded yet")
        bc_actions.append("Upload your business credit report for a complete assessment")

    scores["Your Business Credit"] = {"score": bc_score, "weight": 0.20, "findings": bc_findings, "actions": bc_actions}

    # ── 3. Cash & Runway (20%) ──
    cash_score = 0
    cash_findings = []
    cash_actions = []
    if cash:
        total = cash.get("total_balance", 0)
        accts = cash.get("accounts", [])
        if total >= 500_000: cash_score += 50; cash_findings.append(f"Total cash {fmt(total)} — strong position")
        elif total >= 100_000: cash_score += 35; cash_findings.append(f"Total cash {fmt(total)} — solid")
        elif total >= 25_000: cash_score += 15; cash_findings.append(f"Total cash {fmt(total)} — banks want to see more reserves"); cash_actions.append("Build your cash reserves — banks want to see you can handle a slow month")
        else: cash_findings.append(f"Only {fmt(total)} in cash — this will concern lenders"); cash_actions.append("You need more cash on hand before applying — banks want 3+ months of operating expenses saved")

        if len(accts) >= 2: cash_score += 15; cash_findings.append(f"{len(accts)} bank accounts — shows financial organization")
        elif len(accts) == 1: cash_score += 5; cash_findings.append("Single bank account")

        forecast, _ = run_safe(forecast_cash_position, DATA, 90)
        if forecast and forecast.get("projections"):
            projs = forecast["projections"]
            ending = [p.get("ending_balance", 0) for p in projs]
            if all(e > 0 for e in ending):
                cash_score += 35; cash_findings.append("Cash stays positive for the next 90 days — good sign")
            elif any(e <= 0 for e in ending):
                cash_findings.append("Your cash could run negative in the next 90 days"); cash_actions.append("Fix your cash flow before applying — running out of cash is the #1 reason banks say no")
            else:
                cash_score += 15
    else:
        cash_findings.append("No bank statement data uploaded yet")
        cash_actions.append("Upload your bank statements so we can check your cash position")

    scores["Your Cash Position"] = {"score": cash_score, "weight": 0.20, "findings": cash_findings, "actions": cash_actions}

    # ── 4. Working Capital (15%) ──
    wc_score = 0
    wc_findings = []
    wc_actions = []
    if wc_data:
        cr = wc_data.get("current_ratio", 0)
        dso = wc_data.get("days_sales_outstanding", 0)
        ccc = wc_data.get("cash_conversion_cycle_days", 0)
        runway = wc_data.get("runway_days", 0)

        if cr >= 2.0: wc_score += 35; wc_findings.append(f"Current ratio {cr:.2f}x — you can easily cover your short-term bills")
        elif cr >= 1.5: wc_score += 25; wc_findings.append(f"Current ratio {cr:.2f}x — healthy")
        elif cr >= 1.2: wc_score += 15; wc_findings.append(f"Current ratio {cr:.2f}x — a little tight"); wc_actions.append("Try to get your current ratio above 1.5x by paying down short-term debts")
        elif cr >= 1.0: wc_score += 5; wc_findings.append(f"Current ratio {cr:.2f}x — banks prefer at least 1.25x"); wc_actions.append(f"Your {cr:.2f}x ratio means you're barely covering your bills — build up current assets")
        else: wc_findings.append(f"Current ratio {cr:.2f}x — you owe more short-term than you have"); wc_actions.append("Fix your negative working capital before applying — this will block most financing")

        if dso <= 30: wc_score += 25; wc_findings.append(f"Getting paid in {dso:.0f} days on average — fast collections")
        elif dso <= 45: wc_score += 15; wc_findings.append(f"Getting paid in {dso:.0f} days — reasonable")
        elif dso <= 60: wc_score += 5; wc_findings.append(f"Taking {dso:.0f} days to collect payment — banks see this as slow"); wc_actions.append(f"Speed up your collections — offer 2% discounts for paying within 10 days")
        else: wc_findings.append(f"Taking {dso:.0f} days to collect — this is a problem"); wc_actions.append(f"Your customers are taking {dso:.0f} days to pay — tighten your payment terms now")

        if ccc <= 30: wc_score += 20; wc_findings.append(f"Cash cycle {ccc:.0f} days — money flows back quickly")
        elif ccc <= 60: wc_score += 10; wc_findings.append(f"Cash cycle {ccc:.0f} days — moderate")
        else: wc_findings.append(f"Cash cycle {ccc:.0f} days — your money is tied up too long"); wc_actions.append("Shorten your cash cycle — collect faster and negotiate longer payment terms with suppliers")

        if runway >= 180: wc_score += 20; wc_findings.append(f"{runway:.0f} days of runway — comfortable buffer")
        elif runway >= 90: wc_score += 10; wc_findings.append(f"{runway:.0f} days of runway — adequate")
        else: wc_findings.append(f"Only {runway:.0f} days of runway — banks want 90+ days"); wc_actions.append(f"Build up from {runway:.0f} to at least 90 days of runway before applying")
    else:
        wc_findings.append("Working capital details not entered yet")
        wc_actions.append("Enter your working capital numbers above for a more accurate score")

    scores["Working Capital"] = {"score": wc_score, "weight": 0.15, "findings": wc_findings, "actions": wc_actions}

    # ── 5. Existing Debt (10%) ──
    debt_score = 0
    debt_findings = []
    debt_actions = []
    if cov:
        breaches = cov.get("breaches", 0)
        warnings = cov.get("warnings", 0)
        overall = cov.get("overall_status", "unknown")

        if overall == "compliant" and breaches == 0 and warnings == 0:
            debt_score = 100; debt_findings.append("All existing loan agreements are in good standing")
        elif overall == "compliant" and warnings > 0:
            debt_score = 60; debt_findings.append(f"You're compliant but have {warnings} warning(s) — a new lender will ask about these"); debt_actions.append("Get more headroom on your existing loan covenants before taking on new debt")
        elif breaches > 0:
            debt_findings.append(f"You have {breaches} covenant breach(es) — this will likely block new financing")
            debt_actions.append("Resolve all existing loan violations before applying for anything new")
    else:
        debt_score = 70
        debt_findings.append("No existing loan data on file — that's not necessarily bad")

    scores["Existing Debt"] = {"score": debt_score, "weight": 0.10, "findings": debt_findings, "actions": debt_actions}

    # ── 6. Documentation (10%) ──
    doc_score = 0
    doc_findings = []
    doc_actions = []
    doc_checklist = {
        "Bank statements": cash is not None,
        "Personal credit report": credit is not None and bool(credit.get("personal_profiles")),
        "Business credit report": credit is not None and bool(credit.get("business_profile")),
        "Existing debt schedule": cov is not None,
        "Working capital data": wc_data is not None,
    }
    docs_present = sum(1 for v in doc_checklist.values() if v)
    docs_total = len(doc_checklist)
    doc_score = int((docs_present / docs_total) * 100)
    doc_findings.append(f"{docs_present} of {docs_total} key documents loaded")
    missing = [k for k, v in doc_checklist.items() if not v]
    if missing:
        doc_actions.append(f"Still needed: {', '.join(missing)}")

    scores["Documentation"] = {"score": doc_score, "weight": 0.10, "findings": doc_findings, "actions": doc_actions}

    # ════════════════════════════════════════════════════════
    #  OVERALL SCORE
    # ════════════════════════════════════════════════════════
    overall_score = min(sum(s["score"] * s["weight"] for s in scores.values()), 100)

    if overall_score >= 75: overall_status = "strong"
    elif overall_score >= 55: overall_status = "fair"
    elif overall_score >= 35: overall_status = "needs_work"
    else: overall_status = "not_ready"

    score_color = GREEN if overall_score >= 75 else YELLOW if overall_score >= 50 else RED

    st.markdown(f"""<div class="card" style="text-align:center;padding:2rem;">
<div class="card-header">YOUR FINANCING READINESS SCORE</div>
<div style="font-size:4rem;font-weight:800;color:{score_color};line-height:1;">{overall_score:.0f}</div>
<div style="font-size:1rem;color:{TEXT_SECONDARY};margin-top:0.5rem;">out of 100</div>
<div style="margin-top:1rem;">{tag_html(overall_status)}</div>
</div>""", unsafe_allow_html=True)

    # ── Category Breakdown ──
    st.markdown('<hr class="section-divider">', unsafe_allow_html=True)
    st.markdown("### What's Helping and What's Hurting")

    for cat_name, cat_data in scores.items():
        cat_score = cat_data["score"]
        cat_weight = cat_data["weight"]
        cat_weighted = cat_score * cat_weight
        if cat_score >= 70: cat_status = "strong"
        elif cat_score >= 40: cat_status = "fair"
        else: cat_status = "weak"

        with st.expander(f"{cat_name} — {cat_score}/100 ({cat_weight*100:.0f}% of your score)"):
            c1, c2, c3 = st.columns(3)
            c1.metric("Score", f"{cat_score}/100")
            c2.metric("Weight", f"{cat_weight*100:.0f}%")
            c3.metric("Points", f"{cat_weighted:.0f}")

            if cat_data["findings"]:
                st.markdown("**What we found:**")
                for f in cat_data["findings"]:
                    st.markdown(status_html(cat_status, f), unsafe_allow_html=True)

            if cat_data["actions"]:
                st.markdown("**What to do:**")
                for a in cat_data["actions"]:
                    st.markdown(f"- {a}")

    # ── Top Actions ──
    st.markdown('<hr class="section-divider">', unsafe_allow_html=True)
    st.markdown("### Your Top Action Items")
    st.caption("Fix these before walking into the bank.")

    all_actions = []
    priority_order = ["Your Personal Credit", "Your Business Credit", "Your Cash Position", "Existing Debt", "Working Capital", "Documentation"]
    for cat in priority_order:
        if cat in scores:
            for action in scores[cat]["actions"]:
                all_actions.append({"category": cat, "action": action, "score": scores[cat]["score"]})
    all_actions.sort(key=lambda x: x["score"])

    if all_actions:
        for i, item in enumerate(all_actions[:8], 1):
            urgency = "weak" if item["score"] < 30 else "fair" if item["score"] < 60 else "strong"
            action_text = item["action"]
            cat_text = item["category"]
            st.markdown(status_html(urgency, f"{i}. {action_text}", cat_text), unsafe_allow_html=True)
    else:
        st.markdown(status_html("strong", "You look ready — no major issues found"), unsafe_allow_html=True)

    # ── Loan Product Matcher ──
    st.markdown('<hr class="section-divider">', unsafe_allow_html=True)
    st.markdown("### Loans You May Qualify For")

    fico_score = 0
    if credit and credit.get("personal_profiles"):
        fico_score = credit["personal_profiles"][0].get("credit_score", 0)
    paydex_score = 0
    yrs_in_biz = 0
    if credit and credit.get("business_profile"):
        paydex_score = credit["business_profile"].get("paydex_score", 0)
        yrs_in_biz = credit["business_profile"].get("years_in_business", 0)
    total_cash_val = cash.get("total_balance", 0) if cash else 0
    has_breach = cov.get("breaches", 0) > 0 if cov else False

    products = []
    if fico_score >= 680 and yrs_in_biz >= 2 and not has_breach:
        rate = "Prime + 1-3%" if fico_score >= 720 else "Prime + 2-5%"
        products.append({"name": "Bank Line of Credit", "likelihood": "strong", "rate": rate,
                         "detail": "A revolving credit line for day-to-day expenses. You draw what you need and only pay interest on what you use."})
    elif fico_score >= 640:
        products.append({"name": "Bank Line of Credit", "likelihood": "fair", "rate": "Prime + 3-6%",
                         "detail": "Possible but you'll likely need collateral or a co-signer due to your credit profile."})

    if fico_score >= 640 and yrs_in_biz >= 1:
        products.append({"name": "SBA 7(a) Loan", "likelihood": "strong" if fico_score >= 680 else "fair",
                         "rate": "Prime + 2.25-4.75%",
                         "detail": "Government-backed loan up to $5M with longer repayment terms. Takes 2-3 months to process but offers better rates than conventional loans."})

    if yrs_in_biz >= 0:
        products.append({"name": "SBA Microloan", "likelihood": "strong", "rate": "6-9%",
                         "detail": "Up to $50K from nonprofit lenders. Great if you're newer in business or still building credit."})

    if fico_score >= 600:
        products.append({"name": "Equipment Financing", "likelihood": "strong" if fico_score >= 680 else "fair",
                         "rate": "5-15%",
                         "detail": "Borrow to buy equipment — the equipment itself is your collateral, making this easier to qualify for."})

    if total_cash_val > 0:
        products.append({"name": "Invoice Financing", "likelihood": "fair",
                         "rate": "1-5% per invoice",
                         "detail": "Sell your unpaid invoices for immediate cash. The lender looks at your customers' credit, not yours."})

    if fico_score >= 700 and yrs_in_biz >= 3 and not has_breach:
        products.append({"name": "Term Loan", "likelihood": "strong", "rate": "Prime + 1-4%",
                         "detail": "A fixed amount with a fixed repayment schedule. Best for specific investments like expansion or real estate."})

    if products:
        for prod in products:
            with st.expander(f'{prod["name"]} — Est. Rate: {prod["rate"]}'):
                st.markdown(f'Likelihood: {tag_html(prod["likelihood"])}', unsafe_allow_html=True)
                st.markdown(prod["detail"])
    else:
        st.warning("Upload your credit reports and bank statements so we can match you with loan products.")

    # ── Documentation Checklist ──
    st.markdown('<hr class="section-divider">', unsafe_allow_html=True)
    st.markdown("### What to Bring to the Bank")
    st.caption("Have these ready when you sit down with a banker.")

    banker_docs = [
        ("3 years of business tax returns", "Every lender will ask for these"),
        ("3 years of personal tax returns", "Needed for the personal guaranty all small business loans require"),
        ("Year-to-date profit & loss statement", "Shows how the business is doing this year"),
        ("Current balance sheet", "Shows what you own vs. what you owe"),
        ("Accounts receivable aging", "Who owes you money and how old those invoices are"),
        ("Accounts payable aging", "What you owe and your payment patterns"),
        ("List of all existing debts", "Every loan, line of credit, and lease with balances and monthly payments"),
        ("Personal financial statement", "Your personal assets and debts — the bank will give you their form"),
        ("Business plan or use-of-funds", "What the money is for and how it helps the business"),
        ("6 months of bank statements", "Shows your real cash flow patterns"),
        ("Business licenses and formation docs", "Articles of incorporation, LLC operating agreement, etc."),
    ]

    for doc_name, doc_desc in banker_docs:
        st.markdown(status_html("strong", doc_name, doc_desc), unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════
#  PAGE 2: MY CASH
# ═══════════════════════════════════════════════════════════
elif page == "My Cash":
    st.markdown(f'<div class="top-bar"><div><div class="top-bar-title">My Cash</div><div class="top-bar-sub">Your money — where it is and where it\'s going &middot; {NOW}</div></div></div>', unsafe_allow_html=True)

    cash, cash_e = run_safe(get_cash_position, DATA)

    if cash_e:
        st.error(cash_e)
    elif cash:
        # ── Top line ──
        c1, c2 = st.columns(2)
        c1.metric("Total Cash on Hand", fmt(cash["total_balance"]))
        c2.metric("Number of Accounts", len(cash["accounts"]))

        # ── Accounts table ──
        st.markdown('<hr class="section-divider">', unsafe_allow_html=True)
        st.markdown("### Your Accounts")
        if cash.get("accounts"):
            df = pd.DataFrame(cash["accounts"])
            cols = [c for c in ["bank_name","account_type","currency","balance"] if c in df.columns]
            df = df[cols].copy()
            if "balance" in df.columns:
                df["balance"] = df["balance"].apply(lambda x: f"${x:,.2f}")
            df.columns = [c.replace("_"," ").title() for c in df.columns]
            st.dataframe(df, use_container_width=True, hide_index=True)

        # ── 90-day forecast ──
        st.markdown('<hr class="section-divider">', unsafe_allow_html=True)
        st.markdown("### 90-Day Cash Forecast")
        st.caption("Based on your recent transaction patterns")

        forecast, fc_e = run_safe(forecast_cash_position, DATA, 90)
        if fc_e:
            st.warning("Upload transaction history to see your forecast.")
        elif forecast and forecast.get("projections"):
            projs = forecast["projections"]
            df = pd.DataFrame(projs)
            if "ending_balance" in df.columns and "period" in df.columns:
                chart_df = df.set_index("period")[["ending_balance"]]
                chart_df.columns = ["Projected Balance"]
                st.line_chart(chart_df, color=ACCENT)

            if forecast.get("deficit_periods"):
                st.error(f'Heads up — you could run low on cash during: {", ".join(forecast["deficit_periods"])}')
            if forecast.get("recommendation"):
                st.info(forecast["recommendation"])

            with st.expander("See weekly detail"):
                display = df.copy()
                for col in ["starting_balance","expected_inflows","expected_outflows","net_flow","ending_balance"]:
                    if col in display.columns:
                        display[col] = display[col].apply(lambda x: f"${x:,.0f}")
                display.columns = [c.replace("_"," ").title() for c in display.columns]
                st.dataframe(display, use_container_width=True, hide_index=True)
    else:
        st.info("Upload your bank statements to see your cash position. Switch to **Upload My Files** in the sidebar.")


# ═══════════════════════════════════════════════════════════
#  PAGE 3: MY CREDIT
# ═══════════════════════════════════════════════════════════
elif page == "My Credit":
    st.markdown(f'<div class="top-bar"><div><div class="top-bar-title">My Credit</div><div class="top-bar-sub">What lenders see when they pull your credit &middot; {NOW}</div></div></div>', unsafe_allow_html=True)

    credit, credit_e = run_safe(parse_credit_report, DATA)
    assessment, assess_e = run_safe(assess_credit_position, DATA)

    if credit_e:
        st.error(credit_e)
    elif credit:
        # ── Personal Credit ──
        if credit.get("personal_profiles"):
            st.markdown("### Your Personal Credit")
            for p in credit["personal_profiles"]:
                fico = p["credit_score"]
                if fico >= 750: fico_desc = "Excellent — you'll get the best rates"
                elif fico >= 700: fico_desc = "Good — you qualify for most loans"
                elif fico >= 680: fico_desc = "Fair — expect higher rates"
                elif fico >= 640: fico_desc = "Below average — limited options"
                else: fico_desc = "Needs work — focus on improving before applying"

                c1, c2, c3 = st.columns(3)
                c1.metric("FICO Score", fico)
                c2.metric("Credit Card Usage", f'{p.get("revolving_utilization_pct",0):.0f}%')
                c3.metric("On-Time Payments", f'{p.get("payment_history_pct",0):.0f}%')

                st.markdown(f"**What this means:** {fico_desc}")

                with st.expander("See full details"):
                    st.markdown(f"""<div class="card"><div class="stat-grid">
                        <div class="stat-item"><div class="stat-label">Open Accounts</div><div class="stat-value">{p.get("open_tradelines","—")}</div></div>
                        <div class="stat-item"><div class="stat-label">Late Payments (30/60/90 day)</div><div class="stat-value">{p.get("late_payments_30d",0)} / {p.get("late_payments_60d",0)} / {p.get("late_payments_90d",0)}</div></div>
                        <div class="stat-item"><div class="stat-label">Credit Card Balances</div><div class="stat-value">{fmt(p.get("total_revolving_balance",0))}</div></div>
                        <div class="stat-item"><div class="stat-label">Collections</div><div class="stat-value">{p.get("collections",0)}</div></div>
                        <div class="stat-item"><div class="stat-label">Credit Limit</div><div class="stat-value">{fmt(p.get("total_revolving_limit",0))}</div></div>
                        <div class="stat-item"><div class="stat-label">Bankruptcies</div><div class="stat-value">{p.get("bankruptcies",0)}</div></div>
                        <div class="stat-item"><div class="stat-label">Oldest Account</div><div class="stat-value">{p.get("oldest_account_years","—")} years</div></div>
                        <div class="stat-item"><div class="stat-label">Recent Credit Checks</div><div class="stat-value">{p.get("recent_inquiries_6mo",0)}</div></div>
                    </div></div>""", unsafe_allow_html=True)

        st.markdown('<hr class="section-divider">', unsafe_allow_html=True)

        # ── Business Credit ──
        if credit.get("business_profile"):
            bp = credit["business_profile"]
            st.markdown("### Your Business Credit")

            paydex = bp.get("paydex_score", 0)
            if paydex >= 80: paydex_desc = "Strong — you pay vendors on time or early"
            elif paydex >= 65: paydex_desc = "Okay — room for improvement"
            elif paydex >= 50: paydex_desc = "Slow payer — banks will notice this"
            else: paydex_desc = "Needs work — prioritize paying vendors on time"

            c1, c2, c3 = st.columns(3)
            c1.metric("Paydex Score", paydex)
            c2.metric("Years in Business", bp.get("years_in_business","—"))
            c3.metric("Payment Trend", (bp.get("payment_trend","—") or "—").replace("_"," ").title())

            st.markdown(f"**What this means:** {paydex_desc}")

            with st.expander("See full details"):
                st.markdown(f"""<div class="card"><div class="stat-grid">
                    <div class="stat-item"><div class="stat-label">Business Name</div><div class="stat-value">{bp.get("business_name","—")}</div></div>
                    <div class="stat-item"><div class="stat-label">D&B Rating</div><div class="stat-value">{bp.get("d_and_b_rating","—")}</div></div>
                    <div class="stat-item"><div class="stat-label">Industry</div><div class="stat-value">{bp.get("industry","—")}</div></div>
                    <div class="stat-item"><div class="stat-label">Avg Days Late</div><div class="stat-value">{bp.get("days_beyond_terms_avg","—")}</div></div>
                    <div class="stat-item"><div class="stat-label">Vendor Relationships</div><div class="stat-value">{bp.get("total_trade_experiences","—")}</div></div>
                    <div class="stat-item"><div class="stat-label">Tax Liens</div><div class="stat-value">{bp.get("liens",0)}</div></div>
                    <div class="stat-item"><div class="stat-label">Judgments</div><div class="stat-value">{bp.get("judgments",0)}</div></div>
                    <div class="stat-item"><div class="stat-label">UCC Filings</div><div class="stat-value">{bp.get("ucc_filings",0)}</div></div>
                </div></div>""", unsafe_allow_html=True)

        # ── Overall Assessment ──
        if assessment:
            st.markdown('<hr class="section-divider">', unsafe_allow_html=True)
            st.markdown("### Overall Credit Picture")

            rating = assessment.get("overall_credit_rating", "unknown")
            rating_plain = {
                "strong": "Strong — you're in good shape for financing",
                "fair": "Fair — some areas to improve before applying",
                "marginal": "Marginal — address the issues below first",
                "weak": "Weak — significant work needed before you'll qualify",
            }.get(rating, rating.replace("_"," ").title())

            st.markdown(f'Overall: {tag_html(rating)} — {rating_plain}', unsafe_allow_html=True)

            capacity = assessment.get("lending_capacity_estimate", "—")
            st.markdown(f'<div class="card"><div class="card-header">ESTIMATED BORROWING POWER</div><div class="card-value">{capacity}</div><div class="card-subtitle">Based on your credit profile and financial data</div></div>', unsafe_allow_html=True)

            if assessment.get("risk_factors"):
                st.markdown("### Things Lenders Will Flag")
                for rf in assessment["risk_factors"]:
                    with st.expander(rf["finding"][:80]):
                        st.markdown(f'**Severity:** {tag_html(rf["severity"])}', unsafe_allow_html=True)
                        st.markdown(f'**How it hurts you:** {rf["impact"]}')
                        st.markdown(f'**What to do:** {rf["recommendation"]}')
    else:
        st.info("Upload your credit reports to see your credit picture. Switch to **Upload My Files** in the sidebar.")


# ═══════════════════════════════════════════════════════════
#  PAGE 4: BUSINESS HEALTH
# ═══════════════════════════════════════════════════════════
elif page == "Business Health":
    st.markdown(f'<div class="top-bar"><div><div class="top-bar-title">Business Health Check</div><div class="top-bar-sub">Is your business running efficiently? &middot; {NOW}</div></div></div>', unsafe_allow_html=True)

    st.caption("Enter your numbers from your balance sheet and P&L to see how your business stacks up.")

    c1, c2 = st.columns(2)
    with c1:
        current_assets = st.number_input("Current Assets ($)", min_value=0, value=5_000_000, step=100_000,
                                          help="Cash + accounts receivable + inventory + anything you can convert to cash within a year")
        accounts_receivable = st.number_input("Accounts Receivable ($)", min_value=0, value=2_500_000, step=100_000,
                                               help="Money your customers owe you")
        annual_revenue = st.number_input("Annual Revenue ($)", min_value=0, value=24_000_000, step=1_000_000,
                                          help="Your total sales for the year")
    with c2:
        current_liabilities = st.number_input("Current Liabilities ($)", min_value=0, value=3_000_000, step=100_000,
                                               help="Bills and debts due within a year")
        accounts_payable = st.number_input("Accounts Payable ($)", min_value=0, value=1_500_000, step=100_000,
                                            help="Money you owe your vendors")
        annual_cogs = st.number_input("Annual Cost of Goods Sold ($)", min_value=0, value=18_000_000, step=1_000_000,
                                       help="What it costs to make/deliver what you sell")

    if st.button("Check My Business Health", type="primary"):
        data, err = run_safe(analyze_working_capital,
                             float(current_assets), float(current_liabilities),
                             float(annual_revenue), float(accounts_receivable),
                             float(accounts_payable), float(annual_cogs))
        if err:
            st.error(err)
        elif data:
            # Overall assessment
            assessment_status = data.get("assessment", "unknown")
            plain_assessment = {
                "strong": "Your business is financially healthy",
                "healthy": "Your business looks good — solid fundamentals",
                "moderate": "Some areas need attention",
                "weak": "There are issues to address",
                "critical": "Your business needs immediate financial attention",
            }.get(assessment_status, assessment_status.replace("_"," ").title())

            st.markdown(f'### {plain_assessment} {tag_html(assessment_status)}', unsafe_allow_html=True)
            st.markdown('<hr class="section-divider">', unsafe_allow_html=True)

            # Key metrics with explanations
            c1, c2, c3 = st.columns(3)
            nwc = data.get("net_working_capital", 0)
            cr = data.get("current_ratio", 0)
            runway = data.get("runway_days", 0)

            c1.metric("Money Left After Bills", fmt(nwc))
            c2.metric("Bill Coverage Ratio", f'{cr:.2f}x')
            c3.metric("How Long Cash Lasts", f'{runway:.0f} days')

            # Plain-language interpretation
            if cr >= 2.0:
                st.success("You have more than enough to cover your short-term obligations. Banks love this.")
            elif cr >= 1.5:
                st.success("Healthy coverage — you can comfortably pay your bills.")
            elif cr >= 1.0:
                st.warning("You can cover your bills, but there's not much cushion. Try to build more buffer.")
            else:
                st.error("You owe more short-term than you have available. This needs attention before seeking financing.")

            st.markdown('<hr class="section-divider">', unsafe_allow_html=True)

            c4, c5, c6 = st.columns(3)
            dso = data.get("days_sales_outstanding", 0)
            dpo = data.get("days_payable_outstanding", 0)
            ccc = data.get("cash_conversion_cycle_days", 0)

            c4.metric("Days to Get Paid", f'{dso:.0f}')
            c5.metric("Days to Pay Vendors", f'{dpo:.0f}')
            c6.metric("Cash Cycle", f'{ccc:.0f} days')

            # Cash cycle explanation
            if ccc <= 30:
                st.markdown(f"Your cash cycle is **{ccc:.0f} days** — money comes back to you quickly. That's efficient.")
            elif ccc <= 60:
                st.markdown(f"Your cash cycle is **{ccc:.0f} days** — reasonable, but there's room to speed things up.")
            else:
                st.markdown(f"Your cash cycle is **{ccc:.0f} days** — that's a long time for your money to be tied up. Consider collecting faster or negotiating longer payment terms with vendors.")

            if runway < 90:
                st.warning(f'Your cash only lasts {runway:.0f} days at current burn rate. Banks like to see at least 90 days of runway.')
